# VERSION 8.5: Adds simple password protection for TPT deployment

import streamlit as st
import google.generativeai as genai
import os
import docx
from dotenv import load_dotenv
import io
import markdown2
from xhtml2pdf import pisa
import json

# --- PASSWORD PROTECTION ---
def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store password.
        else:
            st.session_state["password_correct"] = False

    # Create a placeholder for the password secret if running locally
    if "password" not in st.secrets:
        st.secrets["password"] = "SELROCKS2025" # You can change this default password

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Enter Password to Access", type="password", on_change=password_entered, key="password"
        )
        st.info("👆 Note: For your TPT product, you will share the password in your downloadable PDF.")
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Enter Password to Access", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect. Please try again.")
        return False
    else:
        # Password correct.
        return True

# --- INITIAL SETUP ---
load_dotenv()
st.set_page_config(page_title="SEL Integration Agent", page_icon="🧠", layout="wide")

# --- Initialize Session State ---
SESSION_STATE_DEFAULTS = {
    "ai_response": "", "response_title": "", "student_materials": "",
    "differentiation_response": "", "parent_email": "", "scenario": "",
    "conversation_history": [], "training_module": "", "training_scenario": "",
    "training_feedback": "", "check_in_questions": "", "strategy_response": ""
}
for key, default_value in SESSION_STATE_DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = default_value

# --- CONSTANTS ---
EXAMPLE_LESSON = "Subject: Science\nGrade Level: 4th Grade\nTopic: The Water Cycle\nObjective: Students will be able to identify and describe the four main stages of the water cycle."
SAFETY_SETTINGS = {"HARM_CATEGORY_HARASSMENT": "BLOCK_NONE", "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE", "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE", "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"}
GRADE_LEVELS = ["Kindergarten", "1st Grade", "2nd Grade", "3rd Grade", "4th Grade", "5th Grade", "6th Grade", "7th Grade", "8th Grade", "9th Grade", "10th Grade", "11th Grade", "12th Grade"]
SUBJECTS = ["Science", "History", "English Language Arts", "Mathematics", "Art", "Music"]
COMPETENCIES = {
    "Self-Awareness": ["Identifying Emotions", "Self-Perception", "Recognizing Strengths", "Self-Confidence", "Self-Efficacy"],
    "Self-Management": ["Impulse Control", "Stress Management", "Self-Discipline", "Self-Motivation", "Goal-Setting", "Organizational Skills"],
    "Social Awareness": ["Perspective-Taking", "Empathy", "Appreciating Diversity", "Respect for Others"],
    "Relationship Skills": ["Communication", "Social Engagement", "Building Relationships", "Teamwork", "Conflict Resolution"],
    "Responsible Decision-Making": ["Identifying Problems", "Analyzing Situations", "Solving Problems", "Evaluating", "Reflecting", "Ethical Responsibility"]
}
CASEL_COMPETENCIES = list(COMPETENCIES.keys())

# --- HELPER & PROMPT FUNCTIONS (omitted for brevity in this view, but included in the full code block) ---
# ... All your helper and prompt functions go here ...
def read_document(uploaded_file):
    if uploaded_file.name.endswith(".docx"):
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        return "\n".join([para.text for para in doc.paragraphs])
    elif uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    return None
def create_pdf(markdown_text):
    html_text = markdown2.markdown(markdown_text)
    pdf_file = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html_text), dest=pdf_file)
    if pisa_status.err: return None
    pdf_file.seek(0)
    return pdf_file
def create_docx(text):
    doc = docx.Document()
    doc.add_heading('SEL Integration Plan', 0)
    for line in text.split('\n'):
        if line.startswith('### '): doc.add_heading(line.lstrip('### '), level=3)
        elif line.startswith('## '): doc.add_heading(line.lstrip('## '), level=2)
        elif line.startswith('# '): doc.add_heading(line.lstrip('# '), level=1)
        else: doc.add_paragraph(line)
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file
def format_moves(moves):
    if isinstance(moves, list): return ', '.join(moves)
    elif isinstance(moves, str): return moves
    else: return str(moves)
def get_analysis_prompt(lesson_plan_text, standard="", competency="", skill=""):
    focus_instruction = ""
    if competency and skill: focus_instruction = f"The primary focus for SEL integration should be on **{competency}**, with a specific emphasis on the skill of **{skill}**."
    elif competency: focus_instruction = f"The primary focus for SEL integration should be on **{competency}**."
    standard_instruction = ""
    if standard and standard.strip(): standard_instruction = f"Crucially, all suggestions must also align with this educational standard: '{standard.strip()}'."
    return f"You are an SEL instructional coach. Analyze this lesson plan, suggest SEL integration points in Markdown format. {focus_instruction} {standard_instruction}\n\nLesson Plan:\n---\n{lesson_plan_text}\n---"
def get_creation_prompt(grade_level, subject, topic, competency="", skill=""):
    focus_instruction = ""
    if competency and skill: focus_instruction = f"The lesson's primary SEL focus must be on **{competency}**, specifically developing the skill of **{skill}**."
    elif competency: focus_instruction = f"The lesson's primary SEL focus must be on **{competency}**."
    return f"You are a master K-12 curriculum designer. Create a complete, SEL-integrated lesson plan based on the user's request. {focus_instruction}\n**Request Details:**\n- **Grade Level:** {grade_level}\n- **Subject:** {subject}\n- **Lesson Topic:** {topic}\nGenerate a complete lesson plan using the following structure in clear Markdown format: # Lesson Plan: {subject} - {topic}, ## 🎯 Learning Objectives, ## 🔑 Vocabulary, ## 🤔 Essential Questions, ## 🎣 Opening Hook, ## 🧑‍🏫 Direct Instruction, ## ✍️ Guided Practice, ## 💡 Independent Practice, ## ✅ Closing.\nAfter the main lesson plan, add a separate section for \"## 🧠 SEL Integration Strategies\" formatted exactly like the analysis output."
def get_student_materials_prompt(lesson_plan_output):
    return f"You are a practical instructional designer. Based on the lesson plan, create student-facing materials in Markdown with headings: ### 🎟️ Exit Ticket, ### 🗣️ Think-Pair-Share Prompts, and ### ✍️ Journal Starters.\n\nLesson Plan:\n{lesson_plan_output}"
def get_differentiation_prompt(lesson_plan_output):
    return f"You are an expert in instructional differentiation. Based on the lesson plan, provide strategies to support diverse learners in Markdown format with headings: ###  scaffold Support (For Struggling Learners), ### ⬆️ Extension Activities (For Advanced Learners), ### 🌐 English Language Learner (ELL) Support.\n\nLesson Plan:\n---{lesson_plan_output}---"
def get_scenario_prompt(competency, skill, grade_level):
    return f"You are a creative writer. Generate a short, relatable, school-based scenario for a {grade_level} student. The scenario must require them to use the SEL competency of **{competency}**, focusing on the skill of **{skill}**. Present it in the second person ('You are...'), ending with a question. Make it a single paragraph."
def get_feedback_prompt(scenario, history):
    formatted_history = "\n".join([f"- {entry['role']}: {entry['content']}" for entry in history])
    return f"You are a supportive SEL coach. A student is working through this scenario:\n**Scenario:** {scenario}\n**Conversation History:**\n{formatted_history}\nYour task is to ask one or two reflective, Socratic-style questions to help the student think deeper. Keep your response brief and encouraging."
def get_training_prompt(competency):
    sub_competencies = ", ".join(COMPETENCIES[competency])
    return f"You are an expert SEL facilitator. Create a professional development module on **{competency}**.\nUse clear, accessible language and structure your response in Markdown with these sections:\n## 🧠 Understanding {competency}\n## 🛠️ Key Skills to Develop\nFor 3-4 of these skills ({sub_competencies}), create a subsection:\n### Skill: [Name of the Skill]\n* **What It Is:** A simple explanation.\n* **Why It Matters:** Its importance.\n* **Classroom Move:** A practical strategy.\n## 🤔 A Final Reflection"
def get_training_scenario_prompt(competency, training_module_text):
    return f"You are an SEL facilitator. Based on the training for **{competency}**, create a common classroom scenario to help a teacher practice. End with an open-ended question.\n\n**Training Content:**\n---{training_module_text}---\n"
def get_training_feedback_prompt(competency, scenario, teacher_response):
    return f"You are a supportive SEL coach. The teacher is practicing **{competency}**. \n**Scenario:** {scenario}\n**Teacher's Response:** {teacher_response}\n**Your Task:** Provide constructive, encouraging feedback. Affirm a positive aspect and then ask a reflective question to deepen their practice."
def get_check_in_prompt(grade_level, tone):
    return f"You are an expert teacher. Generate 3-4 creative morning check-in questions for a **{grade_level}** class with a **{tone}** tone. Format as a numbered list."
def get_parent_email_prompt(lesson_plan):
    return f"You are a skilled educator. Based on the provided lesson plan, draft a professional, easy-to-understand email from a teacher to parents.\n**Lesson Plan:**\n---\n{lesson_plan}\n---\n**Email Structure:**\n1.  **What We're Learning:** Simply identify the main SEL skill.\n2.  **How We Practiced:** Briefly describe a classroom activity.\n3.  **Connection at Home:** Provide one simple conversation starter or activity for parents."
def get_strategy_prompt(situation):
    return f"You are an expert, quick-thinking SEL coach. A teacher needs immediate help with a classroom situation.\n**The Situation:** \"{situation}\"\n\nYour task is to provide 2-3 **quick, actionable, in-the-moment strategies** the teacher can use right now.\nFor each strategy, provide:\n- **Strategy:** A clear, bolded name for the strategy.\n- **What to Do:** A 1-2 sentence, step-by-step instruction.\n- **Why It Works:** A brief rationale.\n\nKeep the tone calm, direct, and supportive."
def clear_generated_content():
    keys_to_clear = ["ai_response", "response_title", "student_materials", "differentiation_response", "parent_email"]
    for key in keys_to_clear:
        if key in st.session_state: st.session_state[key] = ""

# --- MAIN APP LOGIC ---
if check_password():
    # --- API CONFIGURATION ---
    try:
        api_key = os.environ['GEMINI_API_KEY']
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash-latest', safety_settings=SAFETY_SETTINGS)
    except KeyError:
        st.error("🔴 GEMINI_API_KEY not found. Please make sure you have a .env file with your key.")
        st.stop()
    except Exception as e:
        st.error(f"🔴 An error occurred during API configuration: {e}")
        st.stop()
    
    # --- USER INTERFACE ---
    st.title("🧠 SEL Integration Agent")
    st.markdown("Your AI-powered instructional coach for Social-Emotional Learning.")

    tab_list = ["Analyze Existing Lesson", "Create New Lesson", "🧑‍🎓 Student Scenarios", "👩‍🏫 Teacher SEL Training", "☀️ Morning Check-in", "🆘 Strategy Finder"]
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(tab_list)

    # ... (All the code for your tabs goes here) ...
    with tab1:
        st.header("Analyze an Existing Lesson Plan")
        st.info("Upload or paste a lesson plan. Use the dropdowns to select a specific SEL focus.")
        st.markdown("**Optional: Add a Specific SEL Focus**")
        col1a, col2a = st.columns(2)
        with col1a:
            analyze_competency = st.selectbox("Select a CASEL Competency", options=CASEL_COMPETENCIES, index=None, placeholder="Choose a competency...", key="analyze_comp")
        with col2a:
            if analyze_competency:
                analyze_skill = st.selectbox("Select a Focused Skill", options=COMPETENCIES[analyze_competency], index=None, placeholder="Choose a skill...", key="analyze_skill")
            else:
                analyze_skill = None
                st.selectbox("Select a Focused Skill", options=[], disabled=True, key="disabled_analyze_skill_fix")
        st.markdown("---")
        uploaded_file = st.file_uploader("Upload a .txt or .docx file", type=["txt", "docx"])
        lesson_text_paste = st.text_area("Or paste the full text of your lesson plan here.", height=200)
        standard_input = st.text_area("(Optional) Paste educational standard(s) here.", placeholder="e.g., CCSS.ELA-LITERACY.RL.5.2", height=100)
        if st.button("🚀 Generate SEL Suggestions"):
            lesson_content = ""
            if uploaded_file: lesson_content = read_document(uploaded_file)
            elif lesson_text_paste: lesson_content = lesson_text_paste
            if lesson_content:
                with st.spinner("🤖 Your SEL coach is thinking..."):
                    try:
                        clear_generated_content()
                        prompt = get_analysis_prompt(lesson_content, standard_input, analyze_competency, analyze_skill)
                        response = model.generate_content(prompt)
                        st.session_state.ai_response = response.text
                        st.session_state.response_title = "SEL Integration Suggestions"
                    except Exception as e: st.error(f"Error during generation: {e}")
            else: st.warning("Please upload or paste a lesson plan to begin.")
    with tab2:
        st.header("Create a New, SEL-Integrated Lesson")
        st.info("Fill in the details below to generate a new lesson plan from scratch.")
        st.markdown("**Optional: Add a Specific SEL Focus**")
        col1c, col2c = st.columns(2)
        with col1c:
            create_competency = st.selectbox("Select a CASEL Competency", options=CASEL_COMPETENCIES, index=None, placeholder="Choose a competency...", key="create_comp")
        with col2c:
            if create_competency:
                create_skill = st.selectbox("Select a Focused Skill", options=COMPETENCIES[create_competency], index=None, placeholder="Choose a skill...", key="create_skill")
            else:
                create_skill = None
                st.selectbox("Select a Focused Skill", options=[], disabled=True, key="disabled_create_skill")
        st.markdown("---")
        with st.form("create_form"):
            create_grade = st.selectbox("Grade Level", options=GRADE_LEVELS, index=0)
            create_subject = st.selectbox("Subject", options=SUBJECTS, index=0)
            create_topic = st.text_area("Lesson Topic or Objective", "The causes and effects of the American Revolution.")
            submitted = st.form_submit_button("✨ Create SEL Lesson Plan")
            if submitted:
                with st.spinner("🛠️ Building your new lesson plan..."):
                    try:
                        clear_generated_content()
                        prompt = get_creation_prompt(create_grade, create_subject, create_topic, create_competency, create_skill)
                        response = model.generate_content(prompt)
                        st.session_state.ai_response = response.text
                        st.session_state.response_title = "Your New SEL-Integrated Lesson Plan"
                    except Exception as e: st.error(f"An error occurred: {e}")
    with tab3:
        st.header("Interactive SEL Scenarios")
        st.info("Select a competency and skill to generate a practice scenario for a student.")
        col1b, col2b, col3b = st.columns(3)
        with col1b:
            scenario_competency = st.selectbox("Select a CASEL Competency", options=CASEL_COMPETENCIES, index=3, key="scenario_comp")
        with col2b:
            scenario_skill = st.selectbox("Select a Focused Skill", options=COMPETENCIES[scenario_competency], index=0, key="scenario_skill")
        with col3b:
            scenario_grade = st.selectbox("Select a Grade Level", options=GRADE_LEVELS, key="scenario_grade")
        if st.button("🎬 Generate New Scenario"):
            with st.spinner("Writing a scenario..."):
                try:
                    prompt = get_scenario_prompt(scenario_competency, scenario_skill, scenario_grade)
                    response = model.generate_content(prompt)
                    st.session_state.scenario = response.text
                    st.session_state.conversation_history = []
                except Exception as e: st.error(f"Could not generate a scenario: {e}")
        if st.session_state.scenario:
            st.markdown("---")
            st.markdown(f"**Scenario:** {st.session_state.scenario}")
            for entry in st.session_state.conversation_history:
                if entry['role'] == 'Student': st.markdown(f"> **You:** {entry['content']}")
                else: st.markdown(f"**Coach:** {entry['content']}")
            student_response = st.text_input("What would you do or say?", key="student_response_input")
            if st.button("💬 Submit Response"):
                if student_response:
                    st.session_state.conversation_history.append({"role": "Student", "content": student_response})
                    with st.spinner("Coach is thinking..."):
                        try:
                            feedback_prompt = get_feedback_prompt(st.session_state.scenario, st.session_state.conversation_history)
                            feedback_response = model.generate_content(feedback_prompt)
                            st.session_state.conversation_history.append({"role": "Coach", "content": feedback_response.text})
                            st.rerun()
                        except Exception as e: st.error(f"Could not get feedback: {e}")
    with tab4:
        st.header("👩‍🏫 Teacher SEL Training")
        st.info("Select a competency to begin a professional, in-depth training module.")
        training_competency = st.selectbox("Select a CASEL Competency to learn about", options=CASEL_COMPETENCIES, index=None, placeholder="Choose a competency...", key="training_comp_select")
        if st.button("🎓 Start Training Module"):
            if training_competency:
                with st.spinner("Preparing your training module..."):
                    try:
                        prompt = get_training_prompt(training_competency)
                        response = model.generate_content(prompt)
                        st.session_state.training_module = response.text
                        st.session_state.training_scenario = ""
                        st.session_state.training_feedback = ""
                    except Exception as e: st.error(f"Could not generate the training module: {e}")
            else: st.warning("Please select a competency to begin.")
        if st.session_state.training_module:
            st.markdown("---")
            st.markdown(st.session_state.training_module)
            st.markdown("---")
            st.subheader("🎬 Let's Try It Out")
            if st.button("Generate a Practice Scenario"):
                with st.spinner("Creating a classroom scenario..."):
                    try:
                        prompt = get_training_scenario_prompt(st.session_state.training_comp_select, st.session_state.training_module)
                        response = model.generate_content(prompt)
                        st.session_state.training_scenario = response.text
                        st.session_state.training_feedback = ""
                    except Exception as e: st.error(f"Could not generate the scenario: {e}")
            if st.session_state.training_scenario:
                st.info(st.session_state.training_scenario)
                teacher_response = st.text_area("How would you respond to this scenario?", key="teacher_response_area")
                if st.button("Get Feedback"):
                    if teacher_response:
                        with st.spinner("Your coach is reviewing your response..."):
                            try:
                                prompt = get_training_feedback_prompt(st.session_state.training_comp_select, st.session_state.training_scenario, teacher_response)
                                response = model.generate_content(prompt)
                                st.session_state.training_feedback = response.text
                            except Exception as e: st.error(f"Could not generate feedback: {e}")
                    else: st.warning("Please enter your response above.")
                if st.session_state.training_feedback:
                    st.markdown("---")
                    st.markdown("#### Coach's Feedback")
                    st.success(st.session_state.training_feedback)
    with tab5:
        st.header("☀️ SEL Morning Check-in")
        st.info("Instantly generate creative questions for your morning meeting or class check-in.")
        col1d, col2d = st.columns(2)
        with col1d:
            check_in_grade = st.selectbox("Select a Grade Level", options=GRADE_LEVELS, key="check_in_grade")
        with col2d:
            check_in_tone = st.selectbox("Select a Tone", options=["Calm", "Energetic", "Reflective", "Fun", "Serious"], key="check_in_tone")
        if st.button("❓ Generate Questions"):
            with st.spinner("Coming up with some good questions..."):
                try:
                    prompt = get_check_in_prompt(check_in_grade, check_in_tone)
                    response = model.generate_content(prompt)
                    st.session_state.check_in_questions = response.text
                except Exception as e: st.error(f"Could not generate questions: {e}")
        if st.session_state.check_in_questions:
            st.markdown("---")
            st.markdown(st.session_state.check_in_questions)
    with tab6:
        st.header("🆘 On-Demand Strategy Finder")
        st.info("Describe a classroom situation to get immediate, actionable SEL strategies.")
        situation = st.text_area("Describe the situation in your classroom:", placeholder="e.g., 'Two students are arguing over a shared resource,' or 'My class is very unfocused after lunch.'", height=150)
        if st.button("💡 Find a Strategy"):
            if situation and situation.strip():
                with st.spinner("Finding effective strategies..."):
                    try:
                        prompt = get_strategy_prompt(situation)
                        response = model.generate_content(prompt)
                        st.session_state.strategy_response = response.text
                    except Exception as e: st.error(f"Could not find a strategy: {e}")
            else: st.warning("Please describe the situation to get a strategy.")
        if st.session_state.strategy_response:
            st.markdown("---")
            st.markdown(st.session_state.strategy_response)

    # --- DISPLAY OUTPUT AREA FOR TABS 1 AND 2 ---
    if st.session_state.ai_response:
        st.markdown("---")
        st.header(st.session_state.response_title)
        st.markdown(st.session_state.ai_response)
        st.markdown("---")
        st.subheader("📧 Parent Communication")
        if st.button("Generate Parent Email"):
            with st.spinner("Drafting a parent email..."):
                try:
                    email_prompt = get_parent_email_prompt(st.session_state.ai_response)
                    email_response = model.generate_content(email_prompt)
                    st.session_state.parent_email = email_response.text
                except Exception as e: st.error(f"An error occurred while generating the email: {e}")
        if st.session_state.parent_email:
            st.text_area("Parent Email Draft", value=st.session_state.parent_email, height=300)
        st.markdown("---")
        st.subheader("👩‍🏫 Generate Student-Facing Materials")
        if st.button("Generate Materials"):
            with st.spinner("✍️ Creating student materials..."):
                try:
                    materials_prompt = get_student_materials_prompt(st.session_state.ai_response)
                    materials_response = model.generate_content(materials_prompt)
                    st.session_state.student_materials = materials_response.text
                except Exception as e: st.error(f"An error occurred while generating materials: {e}")
        if st.session_state.student_materials:
            st.markdown(st.session_state.student_materials)
        st.markdown("---")
        st.subheader("🧠 Differentiate This Lesson")
        if st.button("Generate Differentiation Strategies"):
            with st.spinner("💡 Coming up with strategies for diverse learners..."):
                try:
                    diff_prompt = get_differentiation_prompt(st.session_state.ai_response)
                    diff_response = model.generate_content(diff_prompt)
                    st.session_state.differentiation_response = diff_response.text
                except Exception as e: st.error(f"An error occurred while generating differentiation strategies: {e}")
        if st.session_state.differentiation_response:
            st.markdown(st.session_state.differentiation_response)
        st.markdown("---")
        st.subheader("📥 Download Your Plan")
        full_download_text = st.session_state.ai_response
        if st.session_state.parent_email: full_download_text += "\n\n---\n\n# Parent Communication Draft\n\n" + st.session_state.parent_email
        if st.session_state.student_materials: full_download_text += "\n\n---\n\n# Student-Facing Materials\n\n" + st.session_state.student_materials
        if st.session_state.differentiation_response: full_download_text += "\n\n---\n\n# Differentiation Strategies\n\n" + st.session_state.differentiation_response
        pdf_file = create_pdf(full_download_text)
        docx_file = create_docx(full_download_text)
        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button(label="Download as PDF", data=pdf_file, file_name="sel_plan.pdf", mime="application/pdf")
        with dl_col2:
            st.download_button(label="Download as Word Doc (.docx)", data=docx_file, file_name="sel_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")